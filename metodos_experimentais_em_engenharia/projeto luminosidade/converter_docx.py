"""
Conversor Markdown → DOCX com formatação ABNT NBR 14724 + padrão engenharia/metrologia
ESTO017-17 – Métodos Experimentais em Engenharia – Grupo 1
"""

import subprocess, sys
for pkg in ["python-docx"]:
    subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"],
                          stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

# ═══════════════════════════════════════════════════════════════════════════════
# CONFIGURAÇÕES
# ═══════════════════════════════════════════════════════════════════════════════
MD_FILE  = r"c:\Users\julian\Downloads\projeto luminosidade\Relatorio_Luminosidade_STK3311.md"
IMG_DIR  = r"c:\Users\julian\Downloads\projeto luminosidade\graficos"
OUT_FILE = r"c:\Users\julian\Downloads\projeto luminosidade\Relatorio_Luminosidade_STK3311.docx"

IMG_MAP = {
    "Diagrama_Ishikawa.png":         os.path.join(IMG_DIR, "Diagrama_Ishikawa.png"),
    "Fig1_E_vs_d_linear.png":        os.path.join(IMG_DIR, "Fig1_E_vs_d_linear.png"),
    "Fig2_loglog_MMQ.png":           os.path.join(IMG_DIR, "Fig2_loglog_MMQ.png"),
    "Fig3_coef_sensibilidade.png":   os.path.join(IMG_DIR, "Fig3_coef_sensibilidade.png"),
    "Fig4_expoentes_comparacao.png": os.path.join(IMG_DIR, "Fig4_expoentes_comparacao.png"),
    "Fig5_balanco_incertezas.png":   os.path.join(IMG_DIR, "Fig5_balanco_incertezas.png"),
}

# Fonte padrão ABNT
FONTE_CORPO  = "Times New Roman"
FONTE_CODIGO = "Courier New"
TAMANHO_CORPO   = 12
TAMANHO_CODIGO  = 10
TAMANHO_LEGENDA = 10

# ═══════════════════════════════════════════════════════════════════════════════
# FUNÇÕES AUXILIARES
# ═══════════════════════════════════════════════════════════════════════════════

def set_spacing_15(para):
    """Aplica espaçamento entre linhas 1,5 (ABNT)."""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)

def set_spacing_single(para):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

def font_set(run, bold=False, italic=False, size=TAMANHO_CORPO, name=FONTE_CORPO, color=None):
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(size)
    run.font.name  = name
    if color:
        run.font.color.rgb = RGBColor(*color)

def apply_inline(para, text, base_size=TAMANHO_CORPO, base_font=FONTE_CORPO):
    """Processa negrito, itálico e código inline dentro de um parágrafo."""
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            r = para.add_run(part[3:-3]); font_set(r, bold=True, italic=True, size=base_size, name=base_font)
        elif part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2]); font_set(r, bold=True, size=base_size, name=base_font)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = para.add_run(part[1:-1]); font_set(r, italic=True, size=base_size, name=base_font)
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            font_set(r, size=TAMANHO_CODIGO, name=FONTE_CODIGO)
        else:
            r = para.add_run(part); font_set(r, size=base_size, name=base_font)

def clean(text):
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*',     r'\1', text)
    text = re.sub(r'`([^`]+)`',     r'\1', text)
    return text.strip()

def add_hr(doc):
    p = doc.add_paragraph()
    set_spacing_single(p)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'bottom']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '1');    el.set(qn('w:color'), '888888')
        pBdr.append(el)
    pPr.append(pBdr)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top='', bottom='', left='', right='', color='AAAAAA', sz='4'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBdr')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   val or 'single')
        el.set(qn('w:sz'),    sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBdr.append(el)
    tcPr.append(tcBdr)

def add_page_number_footer(doc):
    """Adiciona número de página centralizado no rodapé (ABNT)."""
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run()
    font_set(run, size=10, name=FONTE_CORPO)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = 'PAGE'
    fldChar2  = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def add_figure(doc, img_path, caption_text, fig_width=Inches(5.8)):
    """Imagem centralizada + legenda abaixo (ABNT: Figura N – Descrição. Fonte: ...)"""
    if not os.path.exists(img_path):
        p = doc.add_paragraph(f"[Imagem não encontrada: {os.path.basename(img_path)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing_single(p)
    p.paragraph_format.space_before = Pt(12)
    try:
        p.add_run().add_picture(img_path, width=fig_width)
    except Exception as e:
        p.add_run(f"[Erro: {e}]")
    # Legenda abaixo da figura
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing_single(cap)
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after  = Pt(14)
    txt = caption_text.strip('*').strip()
    run = cap.add_run(txt)
    font_set(run, size=TAMANHO_LEGENDA, name=FONTE_CORPO, italic=True)

def parse_md_table(lines):
    rows, aligns = [], []
    for ln in lines:
        s = ln.strip()
        if not s.startswith('|'): continue
        if re.match(r'^\|[-:\s|]+\|$', s):
            for c in s.strip('|').split('|'):
                c = c.strip()
                if c.startswith(':') and c.endswith(':'): aligns.append('center')
                elif c.endswith(':'): aligns.append('right')
                else: aligns.append('left')
            continue
        rows.append([c.strip() for c in s.strip('|').split('|')])
    return rows, aligns

def add_table_abnt(doc, rows, aligns, caption=''):
    """Tabela ABNT: legenda ACIMA, borda simples, cabeçalho em negrito."""
    ncols = max(len(r) for r in rows) if rows else 1
    rows  = [r + [''] * (ncols - len(r)) for r in rows]

    # Legenda acima (ABNT: Tabela N – Descrição)
    if caption:
        cap = doc.add_paragraph()
        set_spacing_single(cap)
        cap.paragraph_format.space_before = Pt(12)
        cap.paragraph_format.space_after  = Pt(3)
        run = cap.add_run(caption)
        font_set(run, bold=True, size=TAMANHO_LEGENDA, name=FONTE_CORPO)

    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit   = True

    for ri, row in enumerate(rows):
        is_header = (ri == 0)
        for ci, ct in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            al = aligns[ci] if ci < len(aligns) else 'left'
            para.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                              'right':  WD_ALIGN_PARAGRAPH.RIGHT,
                              'left':   WD_ALIGN_PARAGRAPH.LEFT}[al]
            apply_inline(para, ct, base_size=TAMANHO_LEGENDA)
            if is_header:
                set_cell_bg(cell, '1F497D')
                for r in para.runs: r.bold = True; r.font.color.rgb = RGBColor(255,255,255)
            elif ri % 2 == 1:
                set_cell_bg(cell, 'DCE6F1')

    # Nota abaixo da tabela
    nota = doc.add_paragraph()
    set_spacing_single(nota)
    nota.paragraph_format.space_before = Pt(2)
    nota.paragraph_format.space_after  = Pt(10)
    return tbl

# ═══════════════════════════════════════════════════════════════════════════════
# CRIAÇÃO DO DOCUMENTO
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# ── Margens ABNT NBR 14724 ──────────────────────────────────────────────────
for sec in doc.sections:
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(3.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.0)

# ── Estilos dos títulos (ABNT) ──────────────────────────────────────────────
heading_cfg = {
    'Heading 1': (14, True,  True,  (0x1F,0x49,0x7D)),  # MAIÚSCULAS, negrito
    'Heading 2': (12, True,  False, (0x1F,0x49,0x7D)),  # Negrito
    'Heading 3': (12, True,  False, (0x2E,0x74,0xB5)),  # Negrito
    'Heading 4': (12, False, True,  (0x2E,0x74,0xB5)),  # Itálico
}
for style_name, (sz, bold, _, color) in heading_cfg.items():
    try:
        st = doc.styles[style_name]
        st.font.name  = FONTE_CORPO
        st.font.size  = Pt(sz)
        st.font.bold  = bold
        st.font.color.rgb = RGBColor(*color)
        st.paragraph_format.space_before = Pt(18)
        st.paragraph_format.space_after  = Pt(6)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    except: pass

# Estilo Normal
norm = doc.styles['Normal']
norm.font.name = FONTE_CORPO
norm.font.size = Pt(TAMANHO_CORPO)

add_page_number_footer(doc)

# ════════════════════════════════════════════════════════════════════════════
# CAPA (ABNT NBR 14724)
# ════════════════════════════════════════════════════════════════════════════
def capa_par(doc, txt, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, upper=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    set_spacing_single(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    t = txt.upper() if upper else txt
    r = p.add_run(t)
    font_set(r, bold=bold, size=size, name=FONTE_CORPO)
    return p

capa_par(doc, "UNIVERSIDADE FEDERAL DO ABC — UFABC", bold=True, size=14, upper=False, space_before=0, space_after=4)
capa_par(doc, "Centro de Engenharia, Modelagem e Ciências Sociais Aplicadas", bold=False, size=12, space_after=4)
capa_par(doc, "ESTO017-17 – Métodos Experimentais em Engenharia", bold=False, size=12, space_after=60)

capa_par(doc, "GRUPO 1", bold=True, size=14, space_before=0, space_after=6)
integrantes = [
    ("Gabriela Uieda",   "11202320918"),
    ("Giovanna Andrade", "11202020054"),
    ("Giovanna Neves",   "11202520402"),
    ("Julian Carreiro",  "11201811733"),
]
for nome, ra in integrantes:
    capa_par(doc, f"{nome}   RA: {ra}", bold=False, size=12, space_before=0, space_after=3)

capa_par(doc, "", space_before=40)

capa_par(doc,
    "VERIFICAÇÃO EXPERIMENTAL DA LEI DO INVERSO DO QUADRADO DA DISTÂNCIA\n"
    "PARA LUMINOSIDADE COM SENSOR STK3311-X (SITRONIX)",
    bold=True, size=16, upper=False, space_before=0, space_after=60)

capa_par(doc, "Santo André – SP", bold=False, size=12, space_before=60, space_after=3)
capa_par(doc, "2026", bold=False, size=12, space_before=0, space_after=0)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PROCESSAMENTO DO MARKDOWN
# ════════════════════════════════════════════════════════════════════════════
with open(MD_FILE, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# Pula o cabeçalho YAML/markdown até o RESUMO
start_i = 0
for idx, ln in enumerate(lines):
    if ln.strip().upper().startswith('## RESUMO') or ln.strip().startswith('## RESUMO'):
        start_i = idx
        break

table_counter  = [0]
figure_counter = [0]

i = start_i
n = len(lines)

while i < n:
    raw     = lines[i]
    line    = raw.rstrip('\n')
    stripped = line.strip()

    # ── Linha em branco ────────────────────────────────────────────────────
    if stripped == '':
        i += 1
        continue

    # ── Separador horizontal --- ───────────────────────────────────────────
    if re.match(r'^-{3,}$', stripped):
        add_hr(doc)
        i += 1
        continue

    # ── Cabeçalhos # ──────────────────────────────────────────────────────
    if stripped.startswith('#'):
        level = len(stripped) - len(stripped.lstrip('#'))
        text  = stripped.lstrip('#').strip()
        text  = clean(text)
        if level == 1:
            text = text.upper()
        h = doc.add_heading(text, level=min(level, 4))
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        # Garante fonte correta nos runs
        for run in h.runs:
            run.font.name = FONTE_CORPO
            run.font.size = Pt(heading_cfg.get(f'Heading {min(level,4)}', (12,True,False,(0,0,0)))[0])
        i += 1
        continue

    # ── Bloco de código ``` ────────────────────────────────────────────────
    if stripped.startswith('```'):
        code_lines = []
        i += 1
        while i < n and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        i += 1
        if code_lines:
            tbl = doc.add_table(rows=1, cols=1)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            set_cell_bg(cell, 'F5F5F5')
            cell.paragraphs[0].clear()
            for ci, cl in enumerate(code_lines):
                p = cell.paragraphs[0] if ci == 0 else cell.add_paragraph()
                set_spacing_single(p)
                run = p.add_run(cl)
                font_set(run, size=TAMANHO_CODIGO, name=FONTE_CODIGO)
            # espaço após
            sp = doc.add_paragraph()
            set_spacing_single(sp)
            sp.paragraph_format.space_after = Pt(8)
        continue

    # ── Blockquote > ──────────────────────────────────────────────────────
    if stripped.startswith('>'):
        bq_lines = []
        while i < n and lines[i].strip().startswith('>'):
            bq_lines.append(lines[i].strip().lstrip('>').strip())
            i += 1
        full_bq = ' '.join(bq_lines)

        # Detecta placeholder de imagem
        found_img = False
        for key, path in IMG_MAP.items():
            if key in full_bq:
                figure_counter[0] += 1
                # Extrai legenda da linha com '*Figura...'
                caption = ''
                for bl in bq_lines:
                    bl_clean = bl.strip('*').strip()
                    if bl_clean.startswith('Figura') or bl_clean.startswith('figura'):
                        caption = f"Figura {figure_counter[0]} — " + re.sub(r'^Figura\s*\w*\s*[—-]\s*', '', bl_clean)
                if not caption:
                    caption = f"Figura {figure_counter[0]}"
                add_figure(doc, path, caption, fig_width=Inches(5.6))
                found_img = True
                break

        if not found_img:
            # Nota/observação ABNT
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            apply_inline(p, clean(full_bq), base_size=TAMANHO_LEGENDA)
            for run in p.runs:
                run.italic = True
        continue

    # ── Tabela Markdown | ─────────────────────────────────────────────────
    if stripped.startswith('|'):
        tbl_lines = []
        while i < n and lines[i].strip().startswith('|'):
            tbl_lines.append(lines[i]); i += 1
        rows, aligns = parse_md_table(tbl_lines)
        if rows:
            table_counter[0] += 1
            add_table_abnt(doc, rows, aligns, caption=f"Tabela {table_counter[0]}")
        continue

    # ── Lista numerada/marcador ────────────────────────────────────────────
    if re.match(r'^(\s*)([-*+]|\d+\.)\s+', stripped):
        while i < n:
            ls = lines[i].strip()
            if not ls:
                i += 1; break
            m_num = re.match(r'^(\d+)\.\s+(.*)', ls)
            m_bul = re.match(r'^[-*+]\s+(.*)', ls)
            if m_num:
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.left_indent   = Cm(1.25)
                p.paragraph_format.first_line_indent = Cm(-0.63)
                set_spacing_15(p)
                apply_inline(p, clean(m_num.group(2)))
            elif m_bul:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent   = Cm(1.25)
                p.paragraph_format.first_line_indent = Cm(-0.63)
                set_spacing_15(p)
                apply_inline(p, clean(m_bul.group(1)))
            else:
                break
            i += 1
        continue

    # ── *Nota rodapé de tabela (linha começando com *Nota) ────────────────
    if stripped.startswith('*Nota') or stripped.startswith('*nota'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
        set_spacing_single(p)
        apply_inline(p, stripped.strip('*').strip(), base_size=TAMANHO_LEGENDA)
        for run in p.runs: run.italic = True
        i += 1
        continue

    # ── Parágrafo normal ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    set_spacing_15(p)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)  # Recuo ABNT
    apply_inline(p, stripped)
    i += 1

# ════════════════════════════════════════════════════════════════════════════
# SALVA
# ════════════════════════════════════════════════════════════════════════════
doc.save(OUT_FILE)
sz = os.path.getsize(OUT_FILE)
print(f"Documento ABNT salvo: {OUT_FILE}")
print(f"Tamanho: {sz/1024:.0f} KB")
print(f"Figuras inseridas: {figure_counter[0]}")
print(f"Tabelas numeradas: {table_counter[0]}")
